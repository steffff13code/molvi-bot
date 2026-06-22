from __future__ import annotations

import io

from bot.config import BASE_DIR

_FONT_PATH = BASE_DIR / "assets" / "fonts" / "DejaVuSans.ttf"


def make_txt(text: str) -> bytes:
    """Простой текстовый файл в UTF-8."""
    return text.encode("utf-8")


def make_docx(title: str, text: str) -> bytes:
    """DOCX-документ: заголовок + абзацы транскрипта."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    if title:
        doc.add_heading(title, level=1)

    for block in text.split("\n"):
        block = block.rstrip()
        if block:
            doc.add_paragraph(block)
        else:
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_pdf(title: str, text: str) -> bytes:
    """PDF-документ с поддержкой кириллицы (шрифт DejaVuSans)."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.add_font("DejaVu", "", str(_FONT_PATH))

    if title:
        pdf.set_font("DejaVu", "", 16)
        pdf.multi_cell(
            pdf.epw, 10, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT
        )
        pdf.ln(2)

    pdf.set_font("DejaVu", "", 12)
    for block in text.split("\n"):
        block = block.rstrip()
        if block:
            pdf.multi_cell(
                pdf.epw, 7, block, new_x=XPos.LMARGIN, new_y=YPos.NEXT
            )
        else:
            pdf.ln(4)

    out = pdf.output()
    return bytes(out)


def build_export(fmt: str, title: str, text: str) -> tuple[bytes, str]:
    """Возвращает (содержимое файла, имя файла) для формата txt|pdf|docx."""
    safe_title = (title or "transcript").strip().replace("\n", " ")[:40]
    if not safe_title:
        safe_title = "transcript"

    if fmt == "txt":
        return make_txt(text), f"{safe_title}.txt"
    if fmt == "docx":
        return make_docx(title, text), f"{safe_title}.docx"
    if fmt == "pdf":
        return make_pdf(title, text), f"{safe_title}.pdf"
    raise ValueError(f"Unknown export format: {fmt}")
